import re
import pandas as pd
from pathlib import Path
from typing import List
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from src import config

def clean_text(text: str) -> str:
    """
    Limpia y normaliza el texto.
    Elimina espacios en blanco duplicados, saltos de línea innecesarios
    y caracteres no imprimibles.
    """
    if not text:
        return ""
    # Reemplazar múltiples saltos de línea por uno solo
    text = re.sub(r'\n+', '\n', text)
    # Reemplazar múltiples espacios o tabulaciones por un solo espacio
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def load_txt(file_path: Path) -> List[Document]:
    """Carga y procesa un archivo de texto plano TXT."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        cleaned_content = clean_text(content)
        if not cleaned_content:
            return []
        metadata = {"source": file_path.name, "type": "txt"}
        return [Document(page_content=cleaned_content, metadata=metadata)]
    except Exception as e:
        raise ValueError(f"Error al leer el archivo TXT {file_path.name}: {str(e)}")

def load_pdf(file_path: Path) -> List[Document]:
    """Carga y procesa un archivo PDF página por página."""
    documents = []
    try:
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            cleaned_text = clean_text(text)
            if cleaned_text:
                metadata = {
                    "source": file_path.name,
                    "page": i + 1,
                    "type": "pdf"
                }
                documents.append(Document(page_content=cleaned_text, metadata=metadata))
        return documents
    except Exception as e:
        raise ValueError(f"Error al leer el archivo PDF {file_path.name}: {str(e)}")

def load_docx(file_path: Path) -> List[Document]:
    """Carga y procesa un archivo de Microsoft Word (.docx)."""
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # También extraer texto de tablas si existen
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
                
        content = "\n".join(full_text)
        cleaned_content = clean_text(content)
        if not cleaned_content:
            return []
        
        metadata = {"source": file_path.name, "type": "docx"}
        return [Document(page_content=cleaned_content, metadata=metadata)]
    except Exception as e:
        raise ValueError(f"Error al leer el archivo DOCX {file_path.name}: {str(e)}")

def load_csv(file_path: Path) -> List[Document]:
    """
    Carga y procesa un archivo CSV.
    Convierte cada fila en una representación textual estructurada para RAG.
    """
    documents = []
    try:
        # Detectar delimitador leyendo la primera línea
        df = pd.read_csv(file_path)
        filename = file_path.name
        
        for idx, row in df.iterrows():
            parts = []
            for col in df.columns:
                val = row[col]
                if pd.notna(val):
                    parts.append(f"{col}: {val}")
            
            row_content = f"Registro en fila {idx + 1} de {filename}: " + ", ".join(parts)
            metadata = {
                "source": filename,
                "row": idx + 1,
                "type": "csv"
            }
            documents.append(Document(page_content=row_content, metadata=metadata))
        return documents
    except Exception as e:
        raise ValueError(f"Error al leer el archivo CSV {file_path.name}: {str(e)}")

def load_xlsx(file_path: Path) -> List[Document]:
    """
    Carga y procesa un archivo Excel (.xlsx).
    Lee todas las hojas de trabajo y convierte cada fila en texto estructurado.
    """
    documents = []
    try:
        xls = pd.ExcelFile(file_path)
        filename = file_path.name
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            for idx, row in df.iterrows():
                parts = []
                for col in df.columns:
                    val = row[col]
                    if pd.notna(val):
                        parts.append(f"{col}: {val}")
                
                row_content = f"En la hoja '{sheet_name}', registro en fila {idx + 1} de {filename}: " + ", ".join(parts)
                metadata = {
                    "source": filename,
                    "sheet": sheet_name,
                    "row": idx + 1,
                    "type": "xlsx"
                }
                documents.append(Document(page_content=row_content, metadata=metadata))
        return documents
    except Exception as e:
        raise ValueError(f"Error al leer el archivo Excel {file_path.name}: {str(e)}")

def load_document(file_path: Path) -> List[Document]:
    """
    Detecta la extensión del archivo y llama al cargador correspondiente.
    """
    suffix = file_path.suffix.lower()
    if suffix == '.txt':
        return load_txt(file_path)
    elif suffix == '.pdf':
        return load_pdf(file_path)
    elif suffix == '.docx':
        return load_docx(file_path)
    elif suffix == '.csv':
        return load_csv(file_path)
    elif suffix in ['.xlsx', '.xls']:
        return load_xlsx(file_path)
    else:
        raise ValueError(f"Formato de archivo no soportado: {suffix}")

def split_documents(documents: List[Document]) -> List[Document]:
    """
    Divide una lista de Documentos en fragmentos más pequeños (chunks).
    Conserva los metadatos de origen para cada chunk.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        length_function=len
    )
    return splitter.split_documents(documents)

def process_file(file_path: Path) -> List[Document]:
    """
    Carga un archivo, extrae y limpia el texto, y lo divide en chunks.
    Retorna la lista de chunks generados.
    """
    docs = load_document(file_path)
    chunks = split_documents(docs)
    return chunks
